# -*- coding: utf-8 -*-
"""Generate SRS v2.0 Word document from Markdown content."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Global styles
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True

def add_para(text, bold=False, size=None, color=None, align=None, italic=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

# ====== Cover ======
add_para('软件需求规格说明书 (SRS)', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('大白陪你背单词 — 用户端', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_table(['项目', '内容'], [
    ['文档编号', 'SRS-VOCAB-CLIENT-V2.0'],
    ['版本', 'v2.0'],
    ['日期', '2026-06-16'],
    ['状态', '正式版'],
    ['密级', '内部'],
])
doc.add_page_break()

# ====== Revision History ======
doc.add_heading('修订记录', level=1)
add_table(['版本', '日期', '修改内容', '作者'], [
    ['v1.0', '2026-06-16', '基于效果图完成初稿', '—'],
    ['v1.1', '2026-06-16', '补充收藏/统计/会员/离线/API/异常/SM-2', '—'],
    ['v2.0', '2026-06-16', '完整版：确定技术选型(uni-app x)、对齐后台SRS接口、完整学习组件清单、语音评测方案(腾讯云SOE)、全量功能设计', '—'],
])
doc.add_paragraph()

# ====== 1. Introduction ======
doc.add_heading('1. 引言', level=1)

doc.add_heading('1.1 编写目的', level=2)
add_para('本文档定义「大白陪你背单词」用户端的完整软件需求规格，作为产品、设计、开发、测试各角色的统一依据。本版为全量设计，不存在V1裁剪——所有功能必须完整实现。')

doc.add_heading('1.2 适用范围', level=2)
add_para('覆盖微信小程序、Android、iOS、鸿蒙（HarmonyOS NEXT）四端用户端。不包含管理后台。')

doc.add_heading('1.3 术语与缩略语', level=2)
add_table(['术语', '说明'], [
    ['学习组件', '单词学习过程中的最小交互单元，如读单词、拼单词、英译汉等'],
    ['新学', '首次接触新单词的学习环节'],
    ['复习', '对已学单词的间隔重复复习环节'],
    ['直选', '从若干选项中直接选择正确答案的题型'],
    ['全拼', '通过字母键盘逐字母拼写出单词的题型'],
    ['拼读', '结合语音朗读与部分字母提示的拼写题型'],
    ['录音评价', '用户朗读单词后，系统通过语音识别给出发音评分'],
    ['打卡', '每日完成学习任务后的签到行为'],
    ['词库', '基于标签(word_tag)筛选出的单词集合'],
    ['标签', '学段/考试类别维度(如高考、四级、雅思)'],
    ['频次', '单词在教材素材中出现的累计次数(total_frequency)'],
    ['UTS', 'uni-app x的原生编译语言，编译为Android/iOS/HarmonyOS原生代码'],
    ['SOE', 'Smart Oral Evaluation，腾讯云智聆口语评测'],
    ['SM-2', 'SuperMemo 2，间隔重复算法'],
])

doc.add_heading('1.4 参考文献', level=2)
add_para('SM-2 间隔重复算法 (SuperMemo, Piotr Wozniak)')
add_para('腾讯云智聆口语评测 (SOE) API 文档')
add_para('uni-app x 官方文档 (https://doc.dcloud.net.cn/uni-app-x/)')
add_para('uni-app x 鸿蒙开发指南')
add_para('后台 SRS：背单词系统-单词库管理后台_SRS_v2.50.docx')

# ====== 2. Overall Description ======
doc.add_heading('2. 整体描述', level=1)

doc.add_heading('2.1 产品视角', level=2)
add_para('「大白陪你背单词」是一款面向 K12 学生及英语学习者的智能背单词应用。以 IP 角色"大白"为伴学形象，融合语音识别评分、9 种学习组件、间隔重复算法、频次驱动的智能推荐，提供趣味化、科学化的单词学习体验。')
add_para('核心价值主张：', bold=True)
add_para('1. 多维学习：9 种学习组件覆盖听、说、读、写全维度')
add_para('2. 语音评测：基于腾讯云 SOE 的 AI 录音评价，音素级纠错')
add_para('3. 频次驱动推荐：基于后台教材素材累计的单词频次，高频必考词优先学习')
add_para('4. 伴学 IP：大白形象贯穿全应用')
add_para('5. 科学复习：基于 SM-2 算法的间隔重复')
add_para('6. 全端覆盖：一套代码编译为小程序+Android+iOS+鸿蒙')

doc.add_heading('2.2 用户特征', level=2)
add_table(['用户类型', '特征描述', '典型场景'], [
    ['K12 学生', '按年级选词库，有明确学习进度', '每日按计划新学+复习'],
    ['备考用户', '有考试目标(四六级/考研/雅思托福)', '针对性选择标签词库'],
    ['兴趣用户', '无明确考试目标', '利用碎片时间学习'],
])

doc.add_heading('2.3 运行环境', level=2)
add_table(['平台', '要求'], [
    ['微信小程序', '基础库 >= 3.0.0，微信 >= 8.0'],
    ['Android', '最低 Android 8.0(API 26)，目标 Android 14'],
    ['iOS', '最低 iOS 14.0，目标 iOS 17.0'],
    ['鸿蒙', 'HarmonyOS NEXT (API 12+)'],
])

doc.add_heading('2.4 假设与依赖', level=2)
add_para('- 用户设备有稳定网络（首次加载、同步、录音评测时）')
add_para('- 词库数据由管理后台维护，客户端通过 API 获取')
add_para('- 单词发音音频文件托管于 CDN')
add_para('- 语音评测服务由腾讯云 SOE 提供')
add_para('- 前端查询可用单词：WHERE status=1 AND total_frequency > 0')

# ====== 3. Technical Architecture ======
doc.add_heading('3. 技术架构决策', level=1)

doc.add_heading('3.1 跨平台框架：uni-app x', level=2)
add_para('最终决策：uni-app x（Vue 3 + TypeScript + UTS）', bold=True)
add_para('关键决策因素：', bold=True)
add_para('1. 鸿蒙必选：uni-app x 4.61+ 正式支持 HarmonyOS NEXT，编译为 ArkTS 原生代码')
add_para('2. 小程序为入口：uni-app x 的小程序生态最成熟')
add_para('3. Vue 生态：国内开发者 Vue 技能储备丰富')
add_para('4. 腾讯云 SOE 兼容：同为腾讯生态，小程序端 SDK 集成最顺畅')

add_table(['维度', 'uni-app x', 'Taro', 'Flutter'], [
    ['小程序支持', '原生支持，最成熟', '原生支持', '需转换层，体验差'],
    ['Android/iOS', 'UTS编译原生', 'React Native层', '原生渲染'],
    ['鸿蒙支持', '4.61+正式支持', '2025.6开源支持', '不支持'],
    ['语音评测集成', 'JS SDK+原生插件', 'JS SDK', '需原生桥接'],
    ['学习成本', '低(Vue)', '低(React)', '中(Dart)'],
])

doc.add_heading('3.2 语音评测：腾讯云 SOE', level=2)
add_para('最终决策：腾讯云 SOE（主）+ 讯飞 ISE（备用）', bold=True)
add_para('关键决策因素：', bold=True)
add_para('1. 小程序原生 SDK：SOE 是唯一提供微信小程序原生 SDK 的语音评测服务')
add_para('2. 腾讯生态协同：与微信登录、微信支付同属腾讯生态')
add_para('3. 离线能力：提供离线 SDK，满足离线学习场景')
add_para('4. 鸿蒙适配：通过服务端中转实现')

add_table(['评测场景', 'SOE模式', '返回维度', '客户端展示'], [
    ['读单词', '单词评测', 'PronAccuracy/PronFluency', '0-100分数+等级评语'],
    ['读例句', '句子评测', 'PronAccuracy/PronFluency/PronCompletion', '0-100分数+错误音素高亮'],
    ['拼读模式', '单词评测(带部分文本)', '同读单词', '同读单词'],
])

doc.add_heading('3.3 技术栈总览', level=2)
add_para('框架：uni-app x 4.64+ (Vue 3 + TypeScript)')
add_para('原生编译：UTS → Android/iOS/HarmonyOS 原生代码')
add_para('状态管理：Pinia')
add_para('语音评测：腾讯云 SOE SDK')
add_para('支付：微信支付 / Apple IAP / Google Play Billing / 华为支付')

# ====== 4. Functional Requirements ======
doc.add_heading('4. 功能需求', level=1)

# 4.1 User System
doc.add_heading('4.1 用户系统', level=2)

doc.add_heading('4.1.1 微信一键登录 [P0]', level=3)
add_para('FR-AUTH-001', bold=True)
add_para('界面：Logo + 大白形象 + 绿色「微信一键登录」按钮 + 协议文字')
add_para('流程：点击登录 → 微信授权 → 获取openId → 首次创建账号/老用户直接进入首页')
add_para('备选：拒绝授权→提示需要授权；服务器异常→提示重试')

doc.add_heading('4.1.2 手机号登录(App端/鸿蒙端) [P0]', level=3)
add_para('FR-AUTH-002', bold=True)
add_para('界面：手机号输入框(+86) + 验证码输入框(60s倒计时) + 「登录」按钮')
add_para('规则：验证码5分钟有效；60秒内仅可获取一次；手机号与微信可绑定互通')

doc.add_heading('4.1.3 选择年级(标签) [P0]', level=3)
add_para('FR-AUTH-003', bold=True)
add_para('界面：网格卡片排列，选择年级→确认→保存→进入首页')
add_para('年级对应后台word_tag标签：基础教育(学前班~高考16个)、大学考试(四级/六级/考研/专四/专八)、留学考试(雅思/托福/GRE/SAT)')

doc.add_heading('4.1.4 个人信息管理 [P1]', level=3)
add_para('FR-USER-001', bold=True)
add_para('头像/昵称/手机号/年级 可编辑')

doc.add_heading('4.1.5 多端同步 [P0]', level=3)
add_para('FR-USER-002', bold=True)
add_para('同步范围：学习进度/复习排期/收藏/打卡/设置；增量同步，冲突以最新时间为准')

# 4.2 Home
doc.add_heading('4.2 首页与学习入口', level=2)
add_para('FR-HOME-001 首页 [P0]', bold=True)
add_para('顶部：头像+昵称+连续打卡天数 | 中部：词库标签+进度+「开始学习」按钮 | 底部导航：首页|词库|我的')
add_para('已学完态：按钮变灰 | 今日结束：庆祝画面+大白形象')

# 4.3 Wordbook
doc.add_heading('4.3 词库与学习目标', level=2)
add_para('FR-DICT-001 词库练习入口 [P0]', bold=True)
add_para('当前词库卡片+词库列表(按标签分组) | 词库=按标签筛选status=1且total_frequency>0的单词集合')

add_para('FR-DICT-002 设置学习目标 [P0]', bold=True)
add_para('每日新学数量(5/10/15/20/30词) | 复习数量由SM-2自动计算 | 学习模式(标准/听力强化/全能)')

# 4.4 Learning Module
doc.add_heading('4.4 学习模块', level=2)

doc.add_heading('4.4.1 新学流程 [P0]', level=3)
add_para('FR-LEARN-001', bold=True)
add_para('正面：单词+音标+发音 | 点击查看词义→释义+词性+例句 | "不认识"/"认识"按钮')
add_para('新学结束：庆祝+统计(新学/掌握/用时)+「继续复习」/「返回首页」')

doc.add_heading('4.4.2 复习流程 [P0]', level=3)
add_para('FR-LEARN-002', bold=True)
add_para('SM-2算法决定复习顺序和频率 | 算法智能匹配题型 | 答错增加复习频率')
add_para('复习结束：庆祝+统计(复习/正确/错误/正确率)')

doc.add_heading('4.4.3 学习组件概览（9种）', level=3)
add_table(['编号', '组件名称', '交互类型', '测评维度'], [
    ['LC-01', '读单词(含录音评价)', '语音录制+AI评分', '说'],
    ['LC-02', '拼单词', '字母选择/输入', '写'],
    ['LC-03', '读例句', '语音录制+AI评分', '说'],
    ['LC-04', '英译汉(选择题)', '四选一', '读'],
    ['LC-05', '汉译英-直选', '四选一', '读+写'],
    ['LC-06', '汉译英-全拼', '字母输入', '写'],
    ['LC-07', '汉译英-拼读', '字母输入+语音', '写+说'],
    ['LC-08', '听力(词)-直选', '四选一(纯音频)', '听'],
    ['LC-09', '听力(句)-直选', '四选一(纯音频)', '听'],
])
add_para('优先级由SM-2算法决定，具体调度策略后续专项讨论。', italic=True)

# 4.5 Word Detail
doc.add_heading('4.5 单词详情 [P1]', level=2)
add_para('FR-LEARN-003', bold=True)
add_para('单词+音标+发音 | 释义+词性 | 例句(目标词高亮) | 短语+记忆技巧+变形')
add_para('对齐后台word表字段：en_phonetic/us_phonetic/word_audio_uk/us/part_speech/meaning/sentences/phrases/memory_skill/transforms/word_relate/difficulty')

# 4.6 Favorites
doc.add_heading('4.6 收藏与生词本 [P0]', level=2)
add_para('FR-FAV-001', bold=True)
add_para('列表+筛选(掌握程度)+搜索+左滑取消收藏+复习生词')

# 4.7 Statistics
doc.add_heading('4.7 学习记录与统计 [P0]', level=2)
add_para('FR-STAT-001', bold=True)
add_para('日历热力图+学习数据卡片+词库进度+学习时长趋势')

# 4.8 Profile & VIP
doc.add_heading('4.8 我的（个人中心与会员）', level=2)
add_para('FR-MINE-001 个人中心 [P0]', bold=True)
add_para('头像+昵称+VIP标签 | 学习数据卡片 | 功能入口：收藏/统计/会员/个人信息/设置')

add_para('FR-MINE-002 会员中心 [P0]', bold=True)
add_table(['功能', '免费用户', 'VIP会员'], [
    ['每日新学', '最多10词', '不限'],
    ['学习组件', '5种基础', '全部9种'],
    ['语音评测', '每日5次', '不限'],
    ['词库数量', '1个标签', '全部标签'],
    ['学习统计', '基础数据', '高级统计+趋势图'],
    ['离线学习', '不支持', '支持'],
])
add_para('套餐：月卡¥18/季卡¥48/年卡¥128')
add_para('支付：微信支付(小程序)/Apple IAP(iOS)/Google Play(Android)/华为支付(鸿蒙)')
add_para('iOS必须提供「恢复购买」功能')

# 4.9 Settings
doc.add_heading('4.9 设置 [P2]', level=2)
add_para('学习提醒(开关+时间) | 消息通知(开关) | 关于我们 | 意见反馈 | 退出登录')

# 4.10 Notifications
doc.add_heading('4.10 消息与提醒 [P1]', level=2)
add_para('学习提醒推送：小程序→微信服务通知/App→系统推送/鸿蒙→华为推送')

# ====== 5. Non-functional Requirements ======
doc.add_heading('5. 非功能需求', level=1)

doc.add_heading('5.1 性能需求', level=2)
add_table(['指标', '要求'], [
    ['首页加载', '<=2秒(含缓存)'],
    ['学习组件切换', '<=300ms'],
    ['录音评测响应', '<=5秒(正常网络)'],
    ['音频播放延迟', '<=500ms'],
])

doc.add_heading('5.2 安全性需求', level=2)
add_para('HTTPS加密 | 敏感信息脱敏 | Token 7天有效 | 支付服务端二次校验 | 录音不持久化存储')

doc.add_heading('5.3 可靠性需求', level=2)
add_para('学习进度双写(本地+服务端) | 断网自动切离线 | 评测超时自动降级')

doc.add_heading('5.4 兼容性需求', level=2)
add_para('微信>=8.0 | Android 8.0~14 | iOS 14~17 | HarmonyOS NEXT API12+ | 屏幕4.7~6.9寸')

# ====== 6. UI Requirements ======
doc.add_heading('6. 界面需求', level=1)
add_para('主色：绿色#4CAF50 | IP：大白 | 圆角卡片8-16px | 答题微动画 | 庆祝动画')

doc.add_heading('6.1 主要页面清单(30页)', level=2)
add_table(['编号', '页面', '优先级'], [
    ['P01', '登录页', 'P0'], ['P02', '手机号登录', 'P0'], ['P03', '选择年级', 'P0'],
    ['P04', '首页', 'P0'], ['P05', '词库练习', 'P0'], ['P06', '设置学习目标', 'P0'],
    ['P07', '新学页面', 'P0'], ['P08', '复习页面', 'P0'],
    ['P09', '读单词组件', 'P0'], ['P10', '拼单词组件', 'P0'], ['P11', '读例句组件', 'P1'],
    ['P12', '英译汉组件', 'P0'], ['P13', '汉译英-直选', 'P0'], ['P14', '汉译英-全拼', 'P1'],
    ['P15', '汉译英-拼读', 'P1'], ['P16', '听力(词)-直选', 'P1'], ['P17', '听力(句)-直选', 'P1'],
    ['P18', '单词详情', 'P1'], ['P19', '新学结束页', 'P0'], ['P20', '复习结束页', 'P0'],
    ['P21', '我的页面', 'P0'], ['P22', '会员中心', 'P0'], ['P23', '个人信息', 'P1'],
    ['P24', '修改昵称', 'P2'], ['P25', '绑定手机号', 'P1'], ['P26', '收藏/生词本', 'P0'],
    ['P27', '学习统计', 'P0'], ['P28', '设置页', 'P2'], ['P29', '意见反馈', 'P2'],
    ['P30', '今日学习结束页', 'P0'],
])

# ====== 7. Data Requirements ======
doc.add_heading('7. 数据需求与后台接口对齐', level=1)

doc.add_heading('7.1 核心数据模型', level=2)
add_para('Word(单词) — 对齐后台word表：', bold=True)
add_table(['字段', '类型', '客户端用途'], [
    ['word', 'varchar(100)', '展示'],
    ['en_phonetic/us_phonetic', 'varchar(50)', '展示+发音'],
    ['word_audio_uk/us', 'varchar(255)', '播放'],
    ['part_speech', 'varchar(20)', '展示'],
    ['meaning', 'text', '展示+测评'],
    ['sentences', 'json', '展示+测评'],
    ['phrases', 'json', '展示'],
    ['transforms', 'json', '展示'],
    ['memory_skill', 'text', '展示'],
    ['word_relate', 'varchar(100)', '详情页跳转'],
    ['difficulty', 'tinyint', '算法参考'],
    ['stage_label', 'varchar(200)', '词库筛选'],
    ['total_frequency', 'int', '排序+推荐'],
    ['status', 'tinyint', '客户端仅展示status=1且total_frequency>0'],
])

add_para('WordTag(标签) — 对齐后台word_tag表：', bold=True)
add_table(['字段', '类型', '客户端用途'], [
    ['name', 'varchar(50)', '年级选择+词库展示'],
    ['group_name', 'varchar(50)', '筛选分组'],
    ['sort_order', 'int', '展示排序'],
    ['status', 'tinyint', '仅展示启用(status=1)'],
])

add_para('客户端专有模型：', bold=True)
add_para('User | LearningRecord | ReviewSchedule | CheckIn | Favorite')

doc.add_heading('7.2 前端查询条件', level=2)
add_para('WHERE status=1 AND total_frequency>0 (确保只有已上架且有素材支撑的单词才返回)', bold=True)

# ====== 8. Learning Components Detail ======
doc.add_heading('8. 学习组件完整规格', level=1)

# LC-01
doc.add_heading('LC-01 读单词（含录音评价）', level=2)
add_table(['状态', '界面描述', '触发'], [
    ['初始态', '单词+音标+发音+"按住录音"按钮', '进入组件'],
    ['录音态', '红色脉冲动画+"录音中"+播放标准发音', '按下录音'],
    ['录音过短', 'Toast:"录音时间太短"', '<3秒松手'],
    ['正在分析', '"正在分析..."loading', '录音完成'],
    ['分析超时', 'Toast:"分析超时"', '>10秒'],
    ['高分结果', '分数(绿色)+对勾+"太棒了"+继续', '>=80分'],
    ['低分结果', '分数(红色)+"继续加油"+再试一次/继续', '<60分'],
])

# LC-02
doc.add_heading('LC-02 拼单词', level=2)
add_table(['状态', '界面描述', '触发'], [
    ['初始态', '中文释义+字母占位格+打乱字母区', '进入组件'],
    ['拼写中', '点击字母填入+已用变灰+退格', '用户操作'],
    ['拼对态', '字母格变绿+对勾+自动播放发音', '全部正确'],
    ['拼错态', '错误字母变红+显示正确拼写', '有错误'],
])

# LC-03~09 simplified
for lc_id, lc_name, lc_states in [
    ('LC-03', '读例句', '初始态:目标单词(高亮)+例句+翻译+"按住录音" | 结果态:同LC-01句子评测'),
    ('LC-04', '英译汉(选择题)', '题目态:英文单词+4个中文选项 | 答对:变绿+对勾 | 答错:变红+正确项变绿+释义+继续'),
    ('LC-05', '汉译英-直选', '题目态:中文释义+4个英文选项 | 答对:变绿+发音 | 答错:变红+正确项变绿 | 查看答案:单词+音标+释义+例句'),
    ('LC-06', '汉译英-全拼', '题目态:中文释义+QWERTY键盘+首字母提示 | 拼对:变绿+发音 | 拼错:变红+显示正确'),
    ('LC-07', '汉译英-拼读', '题目态:中文释义+部分字母已给+字母选择+发音 | 拼对:变绿 | 拼错:红色+正确字母变绿'),
    ('LC-08', '听力(词)-直选', '题目态:播放按钮+4个英文选项(自动播放) | 答对:变绿 | 答错:变红+正确项变绿'),
    ('LC-09', '听力(句)-直选', '题目态:播放例句+4个中文选项 | 答对:变绿 | 答错:变红+正确项变绿'),
]:
    doc.add_heading(f'{lc_id} {lc_name}', level=2)
    add_para(lc_states)

# ====== 9. SM-2 Algorithm ======
doc.add_heading('9. SM-2 间隔重复算法规格', level=1)

doc.add_heading('9.1 quality评分映射', level=2)
add_table(['quality', '含义', '映射条件'], [
    ['5', '完美', '直选一次答对/录音>=95/拼写一次正确'],
    ['4', '容易', '答对但犹豫>5s/录音80-94/拼写第二次正确'],
    ['3', '正常但吃力', '多次尝试后答对/录音60-79/点"认识"但不确定'],
    ['2', '困难', '答错后记住/录音40-59/点"不认识"'],
    ['1', '很困难', '连续2次答错/录音<40'],
    ['0', '完全不记得', '直接点"不认识"/全拼连续2次完全错误'],
])

doc.add_heading('9.2 核心逻辑', level=2)
add_para('quality>=3: repetitions=0→interval=1天; =1→6天; >=2→interval×easeFactor; easeFactor更新')
add_para('quality<3: repetitions=0, interval=1, easeFactor不变')

# ====== 10. API ======
doc.add_heading('10. API 接口需求', level=1)
api_groups = [
    ('认证', ['/api/auth/wechat-login POST 微信登录', '/api/auth/phone-login POST 手机号登录', '/api/auth/send-sms POST 发送验证码', '/api/auth/refresh-token POST 刷新token']),
    ('用户', ['/api/user/profile GET/PUT 获取/更新资料', '/api/user/bind-phone POST 绑定手机号']),
    ('词库', ['/api/tags GET 标签列表(status=1)', '/api/tags/{tagId}/words GET 标签下可用单词', '/api/tags/{tagId}/stats GET 标签统计']),
    ('学习', ['/api/learn/today-plan GET 今日计划', '/api/learn/new-words GET 新学单词', '/api/learn/review-words GET 复习单词', '/api/learn/record POST 提交记录', '/api/learn/complete POST 完成学习']),
    ('单词', ['/api/words/{wordId} GET 单词详情', '/api/words/{wordId}/audio GET 发音URL']),
    ('收藏', ['/api/favorites GET/POST/DELETE 收藏列表/添加/取消']),
    ('统计', ['/api/stats/overview GET 概览', '/api/stats/calendar GET 日历', '/api/stats/trend GET 趋势', '/api/checkin POST/GET 打卡/连续天数']),
    ('会员', ['/api/vip/status GET 状态', '/api/vip/plans GET 套餐', '/api/vip/purchase POST 购买', '/api/vip/verify POST 校验', '/api/vip/restore POST 恢复购买']),
    ('语音评测', ['/api/speech/init-eval POST 初始化评测', '/api/speech/submit-eval POST 提交评测']),
]
for group_name, apis in api_groups:
    doc.add_heading(f'10.{api_groups.index((group_name,apis))+1} {group_name}模块', level=2)
    for api in apis:
        add_para(f'• {api}')

# ====== 11. Offline ======
doc.add_heading('11. 离线模式规格', level=1)
add_table(['功能', '在线', '离线', '说明'], [
    ['新学/复习', '✅', '✅', '缓存待学+复习单词'],
    ['录音评测', '✅', '⚠️降级', '跳过评测，标记完成'],
    ['选择/拼写题', '✅', '✅', '纯本地交互'],
    ['听力题', '✅', '✅', '缓存音频'],
    ['打卡', '✅', '✅', '联网后同步'],
    ['会员购买', '✅', '❌', '需联网'],
])
add_para('缓存：50个待学单词+全部复习单词+200个音频LRU(约20MB)+静态资源永久缓存 | 超100MB LRU淘汰 | 联网后自动同步')

# ====== 12. Error Handling ======
doc.add_heading('12. 异常与边界处理', level=1)
add_table(['场景', '处理策略'], [
    ['网络中断', '自动切离线+Toast提示'],
    ['评测超时', '10秒降级跳过+记录进度'],
    ['录音权限被拒', '引导弹窗→系统设置'],
    ['音频加载失败', '重试1次→跳过发音'],
    ['支付失败', '保留订单+重试+30分钟超时关闭'],
    ['数据同步冲突', '服务端数据为准'],
    ['Token过期', '自动refresh→失败跳登录'],
    ['存储不足', '空间<50MB提醒清缓存'],
    ['弱网', '超时15秒+自动重试2次'],
    ['重复点击', '500ms防抖'],
    ['学习异常退出', '每题自动保存本地+重进恢复'],
])

# Save
output_path = r'C:\Users\Administrator\WorkBuddy\2026-06-16-10-25-02\docs\SRS-背单词用户端-v2.0.docx'
doc.save(output_path)
print('Done:', output_path)
