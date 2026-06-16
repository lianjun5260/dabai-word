#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""将 SRS Markdown 文档转换为格式精良的 Word 文档"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 配置 ──────────────────────────────────────────
MD_PATH = r"C:\Users\Administrator\WorkBuddy\2026-06-16-10-25-02\docs\SRS-背单词用户端-v1.1.md"
DOCX_PATH = r"C:\Users\Administrator\WorkBuddy\2026-06-16-10-25-02\docs\SRS-背单词用户端-v1.1.docx"

FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_CODE = "Consolas"

# ── 工具函数 ──────────────────────────────────────

def set_run_font(run, name, size_pt, bold=False, color=None):
    """设置 run 的字体属性"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = name
    # 中文字体
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=1.5):
    """设置段落间距和行距"""
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_spacing


def add_styled_text(paragraph, text, font_name=FONT_BODY, size_pt=11,
                    bold=False, color=None):
    """向段落中添加带格式的文本，处理加粗语法 **text**"""
    pattern = re.compile(r'\*\*(.+?)\*\*')
    pos = 0
    for m in pattern.finditer(text):
        # 普通文本
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_run_font(run, font_name, size_pt)
        # 加粗文本
        run = paragraph.add_run(m.group(1))
        set_run_font(run, font_name, size_pt, bold=True, color=color)
        pos = m.end()
    # 剩余普通文本
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, font_name, size_pt)


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_style(table):
    """美化表格：Table Grid + 表头底色"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 设置表格边框 (Table Grid 效果)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    # 表头底色
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "4CAF50")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


def format_cell(cell, text, font_name=FONT_BODY, size_pt=10, bold=False, align=None):
    """格式化单元格内容"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    add_styled_text(p, text, font_name, size_pt, bold)
    # 单元格内边距
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="40" w:type="dxa"/>'
        '  <w:left w:w="80" w:type="dxa"/>'
        '  <w:bottom w:w="40" w:type="dxa"/>'
        '  <w:right w:w="80" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)


# ── Markdown 解析 ────────────────────────────────

def parse_md(filepath):
    """解析 Markdown 文件，返回结构化内容列表"""
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()

    blocks = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip("\n")

        # 空行
        if line.strip() == "":
            i += 1
            continue

        # 水平线 ---
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # 代码块
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            blocks.append(("code", "\n".join(code_lines)))
            continue

        # 标题
        hm = re.match(r'^(#{1,6})\s+(.+)$', line)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            blocks.append(("heading", level, text))
            i += 1
            continue

        # 表格
        if "|" in line and i + 1 < n and re.match(r'^\|[\s\-:|]+\|', lines[i + 1].strip()):
            table_rows = []
            # 读取表头
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(header)
            i += 1  # skip separator line
            i += 1  # move to data rows
            while i < n and "|" in lines[i] and lines[i].strip():
                row_data = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(row_data)
                i += 1
            blocks.append(("table", table_rows))
            continue

        # 普通段落
        blocks.append(("paragraph", line))
        i += 1

    return blocks


# ── 生成 Word 文档 ───────────────────────────────

def build_docx(blocks, output_path):
    doc = Document()

    # 设置默认样式
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    style.paragraph_format.line_spacing = 1.5

    # ── 封面页 ──────────────────────────────────
    # 空行占位
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 0, 0)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("软件需求规格说明书")
    set_run_font(run, FONT_HEADING, 22, bold=True)
    set_paragraph_spacing(p, 0, 6)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("背单词应用 — 用户端")
    set_run_font(run, FONT_HEADING, 16, bold=True)
    set_paragraph_spacing(p, 6, 20)

    # 文档信息表格
    info_data = [
        ["文档编号", "SRS-VOCAB-CLIENT-V1.0"],
        ["版本", "v1.1"],
        ["日期", "2026-06-16"],
        ["状态", "修订稿"],
    ]
    table = doc.add_table(rows=len(info_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (k, v) in enumerate(info_data):
        format_cell(table.rows[idx].cells[0], k, FONT_BODY, 11, bold=True)
        format_cell(table.rows[idx].cells[1], v, FONT_BODY, 11)
    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
    set_table_style(table)

    # 分页
    doc.add_page_break()

    # ── 修订记录 ────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("修订记录")
    set_run_font(run, FONT_HEADING, 16, bold=True)
    set_paragraph_spacing(p, 6, 6)

    rev_data = [
        ["版本", "日期", "修改内容", "作者"],
        ["v1.0", "2026-06-16", "基于效果图完成初稿", "—"],
        ["v1.1", "2026-06-16",
         "补充收藏/生词本、学习统计、会员系统、API接口、离线模式、异常处理、SM-2算法规格",
         "—"],
    ]
    table = doc.add_table(rows=len(rev_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_data in enumerate(rev_data):
        for c_idx, val in enumerate(row_data):
            format_cell(table.rows[r_idx].cells[c_idx], val, FONT_BODY, 10,
                        bold=(r_idx == 0))
    set_table_style(table)

    doc.add_page_break()

    # ── 正文内容 ────────────────────────────────
    heading_level_map = {1: 1, 2: 2, 3: 3, 4: 4, 5: 4, 6: 4}

    for block in blocks:
        btype = block[0]

        if btype == "heading":
            level, text = block[1], block[2]
            # 映射 markdown heading level → Word heading level
            # 文档中 # 是总标题（已在封面），## → Heading 1, ### → Heading 2, ...
            # 但用户要求 ## → Heading 2, ### → Heading 3, #### → Heading 4
            # python-docx 的 add_heading 的 level 参数直接对应 Word 的 Heading N
            # 不过 Word 最多支持 Heading 9，我们把 # → 跳过(封面已处理)
            # ## → level=2, ### → level=3, #### → level=4
            actual_level = level
            if actual_level > 4:
                actual_level = 4

            p = doc.add_heading(level=actual_level)
            # 清除默认文字后手动添加 run 以控制字体
            # add_heading 返回的段落已有 run，我们清空重写
            p.clear()
            add_styled_text(p, text, FONT_HEADING, _heading_size(actual_level), bold=True)
            set_paragraph_spacing(p, 12, 4)

        elif btype == "table":
            rows_data = block[1]
            n_rows = len(rows_data)
            n_cols = len(rows_data[0])
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 设置表格自动适配
            table.autofit = True
            for r_idx, row_data in enumerate(rows_data):
                for c_idx, val in enumerate(row_data):
                    if c_idx < n_cols:
                        format_cell(table.rows[r_idx].cells[c_idx], val, FONT_BODY, 10,
                                    bold=(r_idx == 0))
            set_table_style(table)
            # 表格后空一行
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 0, 0)

        elif btype == "code":
            code_text = block[1]
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 4, 4, 1.0)
            # 灰色背景模拟：通过底纹实现
            for line in code_text.split("\n"):
                run = p.add_run(line + "\n")
                set_run_font(run, FONT_CODE, 9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # 段落底纹
            pPr = p._element.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F5F5F5"/>')
            pPr.append(shd)

        elif btype == "paragraph":
            text = block[1]

            # 处理无序列表项 (- 开头)
            if text.lstrip().startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.clear()
                content = text.lstrip()[2:]
                add_styled_text(p, content, FONT_BODY, 11)
                set_paragraph_spacing(p, 2, 2, 1.5)
                continue

            # 处理有序列表项 (1. 开头)
            ol_match = re.match(r'^(\d+)\.\s+(.+)$', text.lstrip())
            if ol_match:
                p = doc.add_paragraph(style="List Number")
                p.clear()
                add_styled_text(p, ol_match.group(2), FONT_BODY, 11)
                set_paragraph_spacing(p, 2, 2, 1.5)
                continue

            # 普通段落
            p = doc.add_paragraph()
            add_styled_text(p, text, FONT_BODY, 11)
            set_paragraph_spacing(p, 2, 2, 1.5)

    # ── 保存 ────────────────────────────────────
    doc.save(output_path)
    print(f"文档已生成: {output_path}")


def _heading_size(level):
    """标题字号映射"""
    mapping = {1: 16, 2: 14, 3: 12, 4: 11}
    return mapping.get(level, 11)


# ── 主入口 ──────────────────────────────────────
if __name__ == "__main__":
    blocks = parse_md(MD_PATH)
    build_docx(blocks, DOCX_PATH)
