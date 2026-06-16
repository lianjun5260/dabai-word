# -*- coding: utf-8 -*-
"""Generate UI Design Spec Word document from Markdown."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True

def add_para(text, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table(headers, rows):
    col_count = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=col_count, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci in range(col_count):
            val = row[ci] if ci < len(row) else ''
            table.rows[ri+1].cells[ci].text = str(val)
    return table

with open(r'C:\Users\Administrator\WorkBuddy\2026-06-16-10-25-02\docs\UI设计标准与规范-v1.0.md', 'r', encoding='utf-8') as f:
    md = f.read()

lines = md.split('\n')
i = 0
in_code_block = False
in_table = False
table_rows = []
table_headers = []

def flush_table():
    global in_table, table_rows, table_headers
    if table_headers and table_rows:
        add_table(table_headers, table_rows)
    table_headers = []
    table_rows = []
    in_table = False

while i < len(lines):
    line = lines[i]
    
    if line.strip().startswith('```'):
        in_code_block = not in_code_block
        if in_code_block:
            lang = line.strip()[3:].strip()
            if lang:
                add_para(f'[{lang}]', italic=True)
        else:
            pass
        i += 1
        continue
    
    if in_code_block:
        add_para(line, size=10)
        i += 1
        continue
    
    stripped = line.strip()
    if stripped.startswith('|') and stripped.endswith('|'):
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        if all(re.match(r'^[-:]+$', c) for c in cells):
            i += 1
            continue
        if not in_table:
            table_headers = cells
            in_table = True
        else:
            table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            flush_table()
    
    if line.startswith('# ') and not line.startswith('## '):
        i += 1
        continue
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=2)
    elif line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=3)
    elif line.startswith('##### '):
        doc.add_heading(line[6:].strip(), level=4)
    elif stripped.startswith('---'):
        doc.add_paragraph()
    elif stripped.startswith('> '):
        p = add_para(stripped[2:], color=(80, 80, 80))
    elif stripped.startswith('- '):
        text = stripped[2:]
        p = doc.add_paragraph(style='List Bullet')
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for j, part in enumerate(parts):
            run = p.add_run(part)
            if j % 2 == 1:
                run.bold = True
    elif re.match(r'^\d+\.\s', stripped):
        text = re.sub(r'^\d+\.\s', '', stripped)
        p = doc.add_paragraph(style='List Number')
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for j, part in enumerate(parts):
            run = p.add_run(part)
            if j % 2 == 1:
                run.bold = True
    elif stripped == '':
        pass
    else:
        text = stripped
        bold_match = re.match(r'\*\*(.*)\*\*', text)
        if bold_match and re.sub(r'\*\*.*\*\*', '', text).strip() == '':
            add_para(bold_match.group(1), bold=True)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
    
    i += 1

flush_table()

output_path = r'C:\Users\Administrator\WorkBuddy\2026-06-16-10-25-02\docs\UI设计标准与规范-v1.0.docx'
doc.save(output_path)
print(f'Word saved: {output_path}')
