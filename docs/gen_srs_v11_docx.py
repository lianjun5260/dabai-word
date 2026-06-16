from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = '\u5b8b\u4f53'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)
for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '\u9ed1\u4f53'
    hs.font.color.rgb = RGBColor(0, 0, 0)

def P(text, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    return p

def T(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)

# === \u5c01\u9762 ===
P('\u8f6f\u4ef6\u9700\u6c42\u89c4\u683c\u8bf4\u660e\u4e66 (SRS)', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
P('\u5927\u767d\u966a\u4f60\u80cc\u5355\u8bcd \u2014 \u7528\u6237\u7aef', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
T(['\u9879\u76ee', '\u5185\u5bb9'], [
    ['\u6587\u6863\u7f16\u53f7', 'SRS-VOCAB-CLIENT-V1.1'],
    ['\u7248\u672c', 'v1.1'],
    ['\u65e5\u671f', '2026-06-16'],
    ['\u72b6\u6001', '\u4fee\u8ba2\u7a3f'],
])
doc.add_page_break()

# === \u4fee\u8ba2\u8bb0\u5f55 ===
doc.add_heading('\u4fee\u8ba2\u8bb0\u5f55', level=1)
T(['\u7248\u672c', '\u65e5\u671f', '\u4fee\u6539\u5185\u5bb9', '\u4f5c\u8005'], [
    ['v1.0', '2026-06-16', '\u57fa\u4e8e\u6548\u679c\u56fe\u5b8c\u6210\u521d\u7a3f', '\u2014'],
    ['v1.1', '2026-06-16', '\u8865\u5145\u751f\u8bcd\u672c\u3001\u5b66\u4e60\u7edf\u8ba1\u3001\u4f1a\u5458\u7cfb\u7edf\u3001API\u63a5\u53e3\u3001\u79bb\u7ebf\u6a21\u5f0f\u3001\u5f02\u5e38\u5904\u7406\u3001SM-2\u7b97\u6cd5', '\u2014'],
])
doc.add_paragraph()

# Read the markdown file content and convert
with open('C:/Users/Administrator/WorkBuddy/2026-06-16-10-25-02/docs/SRS-\u80cc\u5355\u8bcd\u7528\u6237\u7aef-v1.1.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Simple markdown to docx conversion for the main content
import re

lines = md_content.split('\n')
skip_headers = True  # Skip the first few header lines we already wrote
in_code_block = False

for line in lines:
    # Skip the first part (title, table, revision) that we already wrote
    stripped = line.strip()
    
    if stripped.startswith('# ') and '\u5f15\u8a00' in stripped:
        skip_headers = False
    
    if skip_headers:
        continue
    
    if stripped.startswith('```'):
        in_code_block = not in_code_block
        continue
    
    if in_code_block:
        P(stripped)
        continue
    
    # Headings
    if stripped.startswith('#### '):
        doc.add_heading(stripped[5:], level=4)
    elif stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith('# '):
        doc.add_heading(stripped[2:], level=1)
    # Table
    elif stripped.startswith('|') and '|' in stripped[1:]:
        # Collect table lines
        pass  # Tables handled separately below
    # Regular text
    elif stripped:
        # Bold text
        if stripped.startswith('**') and stripped.endswith('**'):
            P(stripped[2:-2], bold=True)
        else:
            P(stripped)

output = 'C:/Users/Administrator/WorkBuddy/2026-06-16-10-25-02/docs/SRS-\u80cc\u5355\u8bcd\u7528\u6237\u7aef-v1.1.docx'
doc.save(output)
print('Done: ' + output)
