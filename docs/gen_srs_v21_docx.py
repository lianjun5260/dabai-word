# -*- coding: utf-8 -*-
"""Generate SRS v2.1 Word document from Markdown content."""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Global styles
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True

def add_para(text, bold=False, size=None, color=None, align=None, italic=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table(headers, rows):
    # Ensure all rows have same column count as headers
    col_count = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=col_count, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci in range(col_count):
            val = row[ci] if ci < len(row) else ''
            table.rows[ri+1].cells[ci].text = str(val)
    return table

# Read Markdown
with open(r'C:\Users\Administrator\WorkBuddy\2026-06-16-10-25-02\docs\SRS-背单词用户端-v2.0.md', 'r', encoding='utf-8') as f:
    md = f.read()

# Parse and generate Word
lines = md.split('\n')
i = 0
in_code_block = False
in_table = False
table_rows = []
table_headers = []

def flush_table():
    global in_table, table_rows, table_headers
    if table_headers and table_rows:
        add_table(table_headers, table_rows)
    table_headers = []
    table_rows = []
    in_table = False

while i < len(lines):
    line = lines[i]
    
    # Code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            in_code_block = False
            # Add end marker
            add_para('---', italic=True)
        else:
            in_code_block = True
            lang = line.strip()[3:].strip()
            if lang:
                add_para(f'[{lang}]', italic=True)
        i += 1
        continue
    
    if in_code_block:
        add_para(line, size=10)
        i += 1
        continue
    
    # Table rows
    stripped = line.strip()
    if stripped.startswith('|') and stripped.endswith('|'):
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        # Check if separator row
        if all(re.match(r'^[-:]+$', c) for c in cells):
            # Separator, headers already captured
            i += 1
            continue
        if not in_table:
            table_headers = cells
            in_table = True
        else:
            table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            flush_table()
    
    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        # H1 - skip doc title (already in cover)
        i += 1
        continue
    elif line.startswith('## '):
        text = line[3:].strip()
        doc.add_heading(text, level=1)
    elif line.startswith('### '):
        text = line[4:].strip()
        doc.add_heading(text, level=2)
    elif line.startswith('#### '):
        text = line[5:].strip()
        doc.add_heading(text, level=3)
    elif line.startswith('##### '):
        text = line[6:].strip()
        doc.add_heading(text, level=4)
    elif stripped.startswith('---'):
        # Horizontal rule - just add space
        doc.add_paragraph()
    elif stripped.startswith('> '):
        # Blockquote
        text = stripped[2:]
        p = add_para(text, italic=True, color=(80, 80, 80))
    elif stripped.startswith('- '):
        # List item
        text = stripped[2:]
        # Handle bold within list
        p = doc.add_paragraph(style='List Bullet')
        # Parse bold markers
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for j, part in enumerate(parts):
            run = p.add_run(part)
            if j % 2 == 1:
                run.bold = True
    elif re.match(r'^\d+\.\s', stripped):
        # Numbered list
        text = re.sub(r'^\d+\.\s', '', stripped)
        p = doc.add_paragraph(style='List Number')
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for j, part in enumerate(parts):
            run = p.add_run(part)
            if j % 2 == 1:
                run.bold = True
    elif stripped == '':
        pass  # Skip empty lines
    else:
        # Regular paragraph - handle bold
        text = stripped
        # Check if whole line is bold
        bold_match = re.match(r'\*\*(.*)\*\*', text)
        if bold_match and re.sub(r'\*\*.*\*\*', '', text).strip() == '':
            add_para(bold_match.group(1), bold=True)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
    
    i += 1

# Flush any remaining table
flush_table()

output_path = r'C:\Users\Administrator\WorkBuddy\2026-06-16-10-25-02\docs\SRS-背单词用户端-v2.1.docx'
doc.save(output_path)
print(f'Word document saved: {output_path}')
