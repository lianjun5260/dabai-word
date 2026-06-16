from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)

def add_para(text, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)
    return table

# === 封面 ===
add_para('软件需求规格说明书 (SRS)', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('大白陪你背单词 \u2014 用户端', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_table(['项目', '内容'], [
    ['文档编号', 'SRS-VOCAB-CLIENT-V1.0'],
    ['版本', 'v1.0'],
    ['日期', '2026-06-16'],
    ['状态', '初稿'],
])
doc.add_page_break()

# === 修订记录 ===
doc.add_heading('修订记录', level=1)
add_table(['版本', '日期', '修改内容', '作者'], [
    ['v1.0', '2026-06-16', '基于效果图完成初稿', '\u2014'],
])
doc.add_paragraph()

# === 1. 引言 ===
doc.add_heading('1. 引言', level=1)

doc.add_heading('1.1 编写目的', level=2)
add_para('本文档定义\u300c大白陪你背单词\u300d用户端的完整软件需求规格，作为产品、设计、开发、测试各角色的统一依据。')

doc.add_heading('1.2 适用范围', level=2)
add_para('覆盖微信小程序、Android、iOS 三端用户端。不包含管理后台。')

doc.add_heading('1.3 术语与缩略语', level=2)
add_table(['术语', '说明'], [
    ['学习组件', '单词学习过程中的最小交互单元，如读单词、拼单词、英译汉等'],
    ['新学', '首次接触新单词的学习环节'],
    ['复习', '对已学单词的间隔重复复习环节'],
    ['直选', '从若干选项中直接选择正确答案的题型'],
    ['全拼', '通过字母键盘逐字母拼写出单词的题型'],
    ['拼读', '结合语音朗读与字母拼写的混合题型'],
    ['录音评价', '用户朗读单词后，系统通过语音识别给出发音评分'],
    ['打卡', '每日完成学习任务后的签到行为'],
    ['词库', '按年级或考试类型组织的单词集合'],
])

doc.add_heading('1.4 参考文献', level=2)
add_para('SM-2 间隔重复算法 (SuperMemo)')
add_para('\u300a微信小程序设计指南\u300b')
add_para('Material Design 3 规范')
add_para('iOS Human Interface Guidelines')

# === 2. 整体描述 ===
doc.add_heading('2. 整体描述', level=1)

doc.add_heading('2.1 产品视角', level=2)
add_para('\u300c大白陪你背单词\u300d是一款面向 K12 学生及英语学习者的智能背单词应用。以 IP 角色\u201c大白\u201d为伴学形象，融合语音识别评分、多维度学习组件（读、拼、译、听）、间隔重复算法，提供趣味化、科学化的单词学习体验。')
add_para('核心价值主张：', bold=True)
add_para('1. 多维学习：读、拼、译、听四大维度覆盖，不只是看和记')
add_para('2. 语音评测：AI 录音评价，帮助用户纠正发音')
add_para('3. 伴学 IP：大白形象陪伴，学习过程不枯燥')
add_para('4. 科学复习：基于遗忘曲线的间隔重复算法')

doc.add_heading('2.2 用户特征', level=2)
add_table(['用户类型', '特征描述', '典型场景'], [
    ['K12 学生', '按年级选词库，有明确学习进度', '每日按计划新学+复习，打卡坚持'],
    ['备考用户', '有考试目标（四六级、考研等）', '针对性选择词库，强化训练'],
    ['兴趣用户', '无明确考试目标，以提升词汇量为目的', '利用碎片时间学习'],
])

doc.add_heading('2.3 运行环境', level=2)
add_table(['平台', '要求'], [
    ['微信小程序', '基础库 >= 3.0.0，微信 >= 8.0'],
    ['Android', '最低 Android 8.0（API 26），目标 Android 14'],
    ['iOS', '最低 iOS 14.0，目标 iOS 17.0'],
])

doc.add_heading('2.4 设计与实现约束', level=2)
add_para('- 跨平台框架优先考虑 Taro / uni-app')
add_para('- 后端接口 RESTful API，JSON 格式，HTTPS 传输')
add_para('- 语音评测需要调用后端 AI 接口，对网络有一定依赖')
add_para('- 核心学习流程支持离线使用（离线时跳过录音评测）')

doc.add_heading('2.5 假设与依赖', level=2)
add_para('- 用户设备有稳定网络（首次加载、同步、录音评测时）')
add_para('- 词库数据由管理后台维护')
add_para('- 发音音频文件托管于 CDN')
add_para('- 语音评测服务由第三方 AI 引擎提供')

# === 3. 功能需求 ===
doc.add_heading('3. 功能需求', level=1)

doc.add_heading('3.1 用户系统', level=2)

doc.add_heading('3.1.1 登录', level=3)
add_para('FR-AUTH-001 微信一键登录 [P0]', bold=True)
add_para('界面描述（参考效果图：01登录）：')
add_para('- 页面顶部：应用 Logo + \u201c大白陪你背单词\u201d 标题')
add_para('- 页面中部：大白 IP 形象插画')
add_para('- 页面底部：绿色\u300c微信一键登录\u300d按钮（圆角，宽度撑满两侧留边距）')
add_para('- 按钮下方：灰色辅助文字 \u201c登录即同意\u300a用户协议\u300b和\u300a隐私政策\u300b\u201d')
add_para('交互流程：', bold=True)
add_para('1. 用户点击\u300c微信一键登录\u300d')
add_para('2. 调起微信授权弹窗')
add_para('3. 用户确认授权')
add_para('4. 获取 openId 和基本信息（昵称、头像）')
add_para('5. 首次登录自动创建账号 \u2192 跳转选择年级页')
add_para('6. 老用户直接进入首页')
add_para('备选流程：用户拒绝授权 \u2192 提示\u300c需要授权才能使用\u300d；服务器异常 \u2192 提示\u300c登录失败，请重试\u300d')

doc.add_paragraph()
add_para('FR-AUTH-002 手机号登录（App 端）[P1]', bold=True)
add_para('界面描述（参考效果图：02登录-手机号登录）：')
add_para('- 手机号输入框（左侧 +86 区号前缀，右侧清除按钮）')
add_para('- 验证码输入框（右侧\u300c获取验证码\u300d按钮，60s 倒计时）')
add_para('- \u300c登录\u300d按钮（绿色，验证码填满后可点击）')
add_para('业务规则：验证码有效期 5 分钟；同一手机号 60 秒内仅可获取一次；手机号与微信账号可绑定互通')

doc.add_heading('3.1.2 选择年级', level=3)
add_para('FR-AUTH-003 选择年级 [P0]', bold=True)
add_para('界面描述（参考效果图：03选择年级）：')
add_para('- 顶部标题：\u201c选择你的年级\u201d')
add_para('- 年级选项以网格卡片形式排列（每行 2-3 个）')
add_para('- 年级选项：小学（1-6年级）、初中（7-9年级）、高中（10-12年级）、大学')
add_para('- 底部：\u300c确认\u300d按钮')
add_para('交互：点击选择年级 \u2192 卡片高亮（绿色边框/背景）\u2192 点击确认 \u2192 保存并进入首页')

doc.add_heading('3.1.3 个人信息管理', level=3)
add_para('FR-USER-001 编辑个人信息 [P1]', bold=True)
add_para('界面描述（参考效果图：08个人信息、09修改昵称、10绑定手机号）：')
add_para('- 头像行：左侧\u201c头像\u201d + 右侧圆形头像 + 编辑箭头')
add_para('- 昵称行：左侧\u201c昵称\u201d + 右侧昵称文字 + 编辑箭头')
add_para('- 手机号行：左侧\u201c手机号\u201d + 右侧手机号或\u201c去绑定\u201d + 编辑箭头')
add_para('- 年级行：左侧\u201c年级\u201d + 右侧当前年级 + 编辑箭头')
add_para('修改昵称：输入框 + 字数提示（2-20字符）+ \u300c保存\u300d按钮')
add_para('绑定手机号：手机号输入框 + 验证码 + \u300c绑定\u300d按钮')

doc.add_heading('3.1.4 设置', level=3)
add_para('FR-USER-002 设置 [P2]', bold=True)
add_para('设置页列表项：学习提醒（开关）、消息通知（开关）、关于我们、意见反馈、退出登录（红色文字）')
add_para('意见反馈页：多行文本输入 + \u300c提交\u300d按钮')

doc.add_heading('3.2 首页与学习入口', level=2)
add_para('FR-HOME-001 首页 [P0]', bold=True)
add_para('界面描述（参考效果图：04首页、05首页已学完、24今日学习结束）：')
add_para('- 顶部状态栏：用户头像 + 昵称 + 连续打卡天数')
add_para('- 中部学习进度卡片：词库名称、今日新学 X 词 / 复习 X 词、进度条、\u300c开始学习\u300d按钮（绿色）')
add_para('- 已学完状态：按钮变为\u201c今日已学完\u201d（灰色不可点击）')
add_para('- 今日学习结束：庆祝画面 + 大白形象 + \u201c今日学习已完成！\u201d + \u300c继续学习更多\u300d/\u300c返回\u300d按钮')
add_para('底部导航栏：首页 | 词库 | 我的')

doc.add_heading('3.3 词库与学习目标', level=2)
add_para('FR-DICT-001 词库练习入口 [P0]', bold=True)
add_para('界面描述（参考效果图：01词库练习）：')
add_para('- 当前词库信息卡片：词库名称、词汇总量、已学/未学进度')
add_para('- 词库列表：卡片形式，包含封面图/图标 + 词库名称 + 单词数量 + \u201c开始学习\u201d按钮')
add_para('- 支持按年级/考试类型筛选')

doc.add_paragraph()
add_para('FR-DICT-002 设置学习目标 [P0]', bold=True)
add_para('界面描述（参考效果图：02设置学习目标）：')
add_para('- 每日新学数量选择：横向选项卡（5/10/15/20/30词），选中项高亮绿色')
add_para('- 每日复习数量选择')
add_para('- 学习模式选择：标准模式 / 听力模式 / 全能模式')
add_para('- 底部：\u300c确认\u300d按钮')
add_para('默认每日新学 10 词，复习数量由算法自动计算')

doc.add_heading('3.4 学习模块', level=2)

doc.add_heading('3.4.1 新学流程', level=3)
add_para('FR-LEARN-001 新学单词 [P0]', bold=True)
add_para('界面描述（参考效果图：21新学、22新学_点击查看词义、22新学结束）：')
add_para('\u3010正面（默认态）\u3011', bold=True)
add_para('- 顶部：返回 + 进度指示（3/10）+ 关闭')
add_para('- 中部：单词（大字号）+ 音标 + 发音按钮')
add_para('- 底部：\u300c点击查看词义\u300d按钮')
add_para('\u3010背面（查看词义态）\u3011', bold=True)
add_para('- 保留顶部和单词信息')
add_para('- 显示中文释义、词性标签、例句（英文+中文翻译）')
add_para('- 底部按钮：\u300c不认识\u300d（灰色）+ \u300c认识\u300d（绿色）')
add_para('  - 点击\u300c不认识\u300d \u2192 标记未掌握，进入后续学习组件强化')
add_para('  - 点击\u300c认识\u300d \u2192 跳过强化环节，进入下一个词')
add_para('\u3010新学结束页\u3011', bold=True)
add_para('- 庆祝图标 + \u201c新学完成！\u201d')
add_para('- 统计：新学 X 词、掌握 X 词、用时 X 分钟')
add_para('- \u300c继续复习\u300d/ \u300c返回首页\u300d按钮')

doc.add_heading('3.4.2 复习流程', level=3)
add_para('FR-LEARN-002 复习单词 [P0]', bold=True)
add_para('界面描述（参考效果图：23复习结束）：')
add_para('\u3010复习结束页\u3011', bold=True)
add_para('- 庆祝图标 + \u201c复习完成！\u201d')
add_para('- 统计：复习 X 词、正确 X 词、错误 X 词、正确率 X%')
add_para('- \u300c继续学习\u300d/ \u300c返回首页\u300d按钮')
add_para('复习机制：由 SM-2 间隔重复算法决定复习顺序和频率；题型由系统智能匹配；答错/不熟练词增加复习频率')

doc.add_heading('3.4.3 学习组件', level=3)
add_para('学习组件是单词学习过程中的核心交互单元。系统根据学习阶段（新学/复习）和学习模式配置，自动组合多种学习组件形成学习路径。')

doc.add_heading('3.4.3.1 读单词（含录音评价）', level=4)
add_para('FR-COMP-001 读单词 [P0]', bold=True)
add_para('界面描述（参考效果图：01_1~01_7 读单词系列）：')
add_para('\u3010初始态\u3011 顶部进度+关闭 | 单词+音标+发音 | 底部\u300c按住录音\u300d按钮（灰色圆形麦克风）')
add_para('\u3010录音态\u3011 按住按钮 \u2192 红色脉冲动画 + \u201c录音中\u201d + 同时播放标准发音供跟读')
add_para('\u3010录音少于3秒\u3011 Toast 提示：\u201c录音时间太短，请再试一次\u201d \u2192 回到初始态')
add_para('\u3010正在分析态\u3011 \u201c正在分析...\u201d loading 动画，等待后端评测接口返回')
add_para('\u3010分析超时\u3011 Toast 提示：\u201c分析超时，请重试\u201d \u2192 回到初始态')
add_para('\u3010评价结果-高分\u3011 分数（绿色大字号，如90）+ \u201c太棒了！\u201d + 对勾 + \u300c继续\u300d按钮')
add_para('\u3010评价结果-低分\u3011 分数（橙色/红色，如55）+ \u201c继续加油\u201d + 标准音标提示 + \u300c再试一次\u300d/\u300c继续\u300d按钮')
add_para('业务规则：录音最短 3 秒；评测超时 10 秒；评分 0-100；高分>=80，低分<60；离线跳过评测')

doc.add_heading('3.4.3.2 拼单词', level=4)
add_para('FR-COMP-002 拼单词 [P0]', bold=True)
add_para('界面描述（参考效果图：02/03/04/07 拼单词系列）：')
add_para('\u3010初始态\u3011 中文释义 | 字母占位格（空白待填）| 底部打乱字母选择区')
add_para('\u3010拼写中\u3011 点击字母 \u2192 填入空格 \u2192 已用字母变灰；支持退格')
add_para('\u3010拼对态\u3011 字母格变绿 + 对勾 + 正确提示音 + 自动进入下一题')
add_para('\u3010拼错态\u3011 错误字母变红 + 叉号 + 显示正确拼写 \u2192 允许重试或跳过')
add_para('业务规则：字母区仅提供所需字母；拼对后自动播放发音；连续拼错2次显示正确答案')

doc.add_heading('3.4.3.3 读例句', level=4)
add_para('FR-COMP-003 读例句 [P1]', bold=True)
add_para('界面描述（参考效果图：05读例句）：')
add_para('- 目标单词（高亮）+ 完整例句 + 中文翻译 + 发音按钮')
add_para('- \u300c按住录音\u300d按钮 \u2192 同读单词的录音评测流程')

doc.add_heading('3.4.3.4 英译汉（选择题）', level=4)
add_para('FR-COMP-004 英译汉 [P0]', bold=True)
add_para('界面描述（参考效果图：08_1~08_2 英译汉系列、12~15 听力系列）：')
add_para('\u3010题目态\u3011 英文单词/句子（大字号）+ 4个中文选项（竖向卡片）')
add_para('\u3010答对态\u3011 选中项变绿 + 对勾 + 正确提示音 + 自动下一题')
add_para('\u3010答错态+显示正确释义\u3011 选中项变红 + 叉号 | 正确选项变绿 + 对勾 | 显示释义 + \u300c继续\u300d按钮')
add_para('听力模式：仅播放音频不显示文字，根据听到的选择中文翻译')

doc.add_heading('3.4.3.5 汉译英', level=4)
add_para('汉译英有三种题型：直选、全拼、拼读。')

doc.add_paragraph()
add_para('FR-COMP-005A 汉译英-直选 [P0]', bold=True)
add_para('界面描述（参考效果图：03~06 中译英(词)-直选系列）：')
add_para('\u3010题目态\u3011 中文释义 + 4个英文单词选项')
add_para('\u3010答对态\u3011 选中项变绿 + 对勾')
add_para('\u3010答错态\u3011 选中项变红 + 叉号 | 正确选项变绿')
add_para('\u3010查看答案态\u3011 正确单词+音标+发音+完整释义+例句 + \u300c继续\u300d按钮')

doc.add_paragraph()
add_para('FR-COMP-005B 汉译英-全拼 [P1]', bold=True)
add_para('界面描述（参考效果图：16~18 中译英(词)-全拼系列）：')
add_para('\u3010题目态\u3011 中文释义 + 字母输入区（QWERTY键盘或字母选择）+ 首字母提示')
add_para('\u3010拼对态\u3011 输入区变绿 + 对勾 + 播放发音')
add_para('\u3010拼错态\u3011 输入区变红 + 叉号 + 显示正确拼写 + \u300c继续\u300d按钮')

doc.add_paragraph()
add_para('FR-COMP-005C 汉译英-拼读 [P1]', bold=True)
add_para('界面描述（参考效果图：19~20_2 中译英(词)-拼读系列）：')
add_para('\u3010题目态\u3011 中文释义 + 部分字母已给出（如 a _ _ _ _ _）+ 字母选择区 + 发音按钮')
add_para('\u3010拼对态\u3011 全部正确 \u2192 变绿 + 对勾 + \u201c太棒了！\u201d')
add_para('\u3010拼错态\u3011 错误字母变红 + 正确字母变绿 + \u201c再试一次\u201d 或显示正确答案')

doc.add_heading('3.4.3.6 听力（词/句直选）', level=4)
add_para('FR-COMP-006 听力选择 [P1]', bold=True)
add_para('界面描述（参考效果图：12~15 听力系列）：')
add_para('\u3010听力(词)-直选\u3011 播放按钮 + \u201c点击播放\u201d + 4个英文单词选项')
add_para('\u3010听力(句)-直选\u3011 播放按钮 + \u201c点击播放\u201d + 4个中文翻译选项')
add_para('答对：选中项变绿+对勾；答错：选中项变红，正确项变绿')
add_para('业务规则：每题自动播放一次可重播；不显示文字内容；词题选英文，句题选中文翻译')

doc.add_heading('3.4.4 单词详情', level=3)
add_para('FR-LEARN-003 单词详情 [P1]', bold=True)
add_para('界面描述（参考效果图：07单词详情）：')
add_para('- 顶部：返回 + 标题\u300c单词详情\u300d + 收藏按钮')
add_para('- 单词区域：英文单词（大字号）+ 音标（英式/美式切换）+ 发音按钮')
add_para('- 释义区域：词性标签 + 中文释义（多条分行）')
add_para('- 例句区域：英文例句（目标单词高亮）+ 中文翻译 + 发音按钮')
add_para('- 短语搭配（可选）+ 词根词缀（可选）')

doc.add_heading('3.5 我的（个人中心与会员）', level=2)
add_para('FR-MINE-001 个人中心 [P0]', bold=True)
add_para('界面描述（参考效果图：07我的、07我的-已开通会员）：')
add_para('- 顶部：头像 + 昵称 + 会员状态标签')
add_para('- 学习数据卡片：累计学习天数 / 累计学习单词数 / 连续打卡天数')
add_para('- 功能入口：我的收藏 | 学习记录 | 会员中心 | 个人信息 | 设置')
add_para('未开通会员：会员中心入口显示\u201c开通会员\u201d')
add_para('已开通会员：头像旁 VIP 图标 + 会员中心显示\u201cVIP会员\u00b7到期时间\u201d')

doc.add_paragraph()
add_para('FR-MINE-002 会员中心 [P1]', bold=True)
add_para('展示会员套餐（月卡/季卡/年卡）、权益说明、购买流程、续费入口')

doc.add_heading('3.6 消息与提醒', level=2)
add_para('FR-MSG-001 学习提醒 [P2]', bold=True)
add_para('可在设置中开关学习提醒，自定义提醒时间，到时推送：\u201c该背单词啦！今天还有 X 个单词待学习\u201d')

# === 4. 非功能需求 ===
doc.add_heading('4. 非功能需求', level=1)

doc.add_heading('4.1 性能需求', level=2)
add_table(['指标', '要求'], [
    ['首页加载时间', '<= 2 秒（含本地缓存）'],
    ['学习组件切换', '<= 300ms'],
    ['录音评测响应', '<= 5 秒（正常网络）'],
    ['音频播放延迟', '<= 500ms'],
    ['离线学习可用', '核心学习组件离线可用（录音评测除外）'],
])

doc.add_heading('4.2 可用性需求', level=2)
add_para('- 学习流程操作步数最简化，核心路径 \u2264 3 步进入学习')
add_para('- 按钮尺寸 >= 44px（符合移动端触控标准）')
add_para('- 文字清晰可读，字号 >= 14px（正文）')
add_para('- 色彩对比度符合 WCAG AA 标准')

doc.add_heading('4.3 安全性需求', level=2)
add_para('- 所有 API 请求使用 HTTPS 加密')
add_para('- 用户敏感信息（手机号）脱敏展示')
add_para('- 登录态 token 有效期 7 天，支持静默续期')
add_para('- 语音数据不上传至第三方（仅发送至评测引擎处理后销毁）')

doc.add_heading('4.4 可靠性需求', level=2)
add_para('- 学习进度实时保存，异常退出不丢失')
add_para('- 网络中断时自动切换离线模式')
add_para('- 语音评测接口超时自动降级（跳过评测环节）')

doc.add_heading('4.5 兼容性需求', level=2)
add_para('- 微信小程序兼容 iOS / Android 微信客户端')
add_para('- Android 8.0 ~ Android 14 全覆盖')
add_para('- iOS 14.0 ~ iOS 17.0 全覆盖')
add_para('- 适配屏幕尺寸：4.7 寸 ~ 6.9 寸')

# === 5. 界面需求 ===
doc.add_heading('5. 界面需求', level=1)

doc.add_heading('5.1 设计原则与视觉风格', level=2)
add_para('基于效果图归纳的设计风格：', bold=True)
add_para('- 主色调：绿色系（#4CAF50 及变体），传达学习、成长、积极的品牌感')
add_para('- 辅助色：白色背景、灰色文字、红色（错误态）、橙色（低分提示）')
add_para('- IP 形象：大白角色贯穿全应用，出现在登录页、首页、学习结束页等')
add_para('- 圆角卡片：界面元素以圆角卡片为主（8-16px 圆角），柔和亲切')
add_para('- 动画：答题反馈（对/错）有微动画，学习结束有庆祝动画')

doc.add_heading('5.2 主要页面清单', level=2)
add_table(['编号', '页面名称', '入口', '优先级'], [
    ['P01', '登录页', '启动/未登录', 'P0'],
    ['P02', '手机号登录页', '登录页切换', 'P1'],
    ['P03', '选择年级页', '首次登录', 'P0'],
    ['P04', '首页', '底部导航', 'P0'],
    ['P05', '词库练习页', '底部导航', 'P0'],
    ['P06', '设置学习目标页', '词库练习页', 'P0'],
    ['P07', '新学页面', '首页-开始学习', 'P0'],
    ['P08', '复习页面', '新学完成后', 'P0'],
    ['P09', '读单词组件', '学习流程中', 'P0'],
    ['P10', '拼单词组件', '学习流程中', 'P0'],
    ['P11', '读例句组件', '学习流程中', 'P1'],
    ['P12', '英译汉组件', '学习流程中', 'P0'],
    ['P13', '汉译英-直选', '学习流程中', 'P0'],
    ['P14', '汉译英-全拼', '学习流程中', 'P1'],
    ['P15', '汉译英-拼读', '学习流程中', 'P1'],
    ['P16', '听力选择组件', '学习流程中', 'P1'],
    ['P17', '单词详情页', '学习中/收藏', 'P1'],
    ['P18', '新学结束页', '新学完成', 'P0'],
    ['P19', '复习结束页', '复习完成', 'P0'],
    ['P20', '我的页面', '底部导航', 'P0'],
    ['P21', '会员中心', '我的页面', 'P1'],
    ['P22', '个人信息页', '我的页面', 'P1'],
    ['P23', '修改昵称', '个人信息页', 'P2'],
    ['P24', '绑定手机号', '个人信息页', 'P1'],
    ['P25', '设置页', '我的页面', 'P2'],
    ['P26', '意见反馈页', '设置页', 'P2'],
])

doc.add_heading('5.3 页面流转关系', level=2)
add_para('启动 \u2192 登录页 \u2192 微信授权/手机号登录 \u2192 选择年级(首次) \u2192 首页')
add_para('首页 \u2192 开始学习 \u2192 新学 \u2192 学习组件轮换(读/拼/译/听) \u2192 新学结束 \u2192 复习 \u2192 复习结束 \u2192 首页')
add_para('首页 \u2192 词库 \u2192 设置目标 \u2192 新学')
add_para('首页 \u2192 我的 \u2192 个人信息/会员中心/设置')

# === 6. 数据需求 ===
doc.add_heading('6. 数据需求', level=1)

doc.add_heading('6.1 核心数据模型', level=2)
add_para('用户(User)：userId, openId, phone, nickname, avatar, grade, vipStatus, vipExpireTime, dailyGoal, remindTime, createdAt, updatedAt')
add_para('词库(Wordbook)：bookId, name, grade, coverUrl, totalWords, description')
add_para('单词(Word)：wordId, bookId, word, phoneticUK, phoneticUS, audioUrl, pos, meanings[], examples[], phrases[], roots')
add_para('学习记录(LearningRecord)：recordId, userId, wordId, learnType, componentType, score, isCorrect, timestamp')
add_para('复习排期(ReviewSchedule)：scheduleId, userId, wordId, nextReviewAt, interval, easeFactor, repetitions')
add_para('打卡记录(CheckIn)：userId, date, wordsLearned, wordsReviewed, streakDays')
add_para('收藏(Favorite)：userId, wordId, createdAt')

doc.add_heading('6.2 数据字典', level=2)
add_table(['字段', '类型', '说明'], [
    ['userId', 'String', '用户唯一标识，UUID'],
    ['openId', 'String', '微信 openId'],
    ['phone', 'String', '手机号（加密存储）'],
    ['grade', 'Int', '年级代码（1-12小学到高三，13+大学）'],
    ['vipStatus', 'Int', '0=免费, 1=会员'],
    ['dailyGoal', 'Int', '每日新学目标数'],
    ['bookId', 'String', '词库 ID'],
    ['wordId', 'String', '单词 ID'],
    ['learnType', 'Enum', 'NEW=新学, REVIEW=复习'],
    ['componentType', 'Enum', 'READ=读, SPELL=拼, EN2CN=英译汉, CN2EN=汉译英, LISTEN=听'],
    ['score', 'Int', '录音评测分数 0-100'],
    ['isCorrect', 'Boolean', '答题是否正确'],
    ['nextReviewAt', 'DateTime', '下次复习时间'],
    ['interval', 'Int', '复习间隔（天）'],
    ['easeFactor', 'Float', 'SM-2 难度因子'],
    ['streakDays', 'Int', '连续打卡天数'],
])

output = 'C:/Users/Administrator/WorkBuddy/2026-06-16-10-25-02/docs/SRS-背单词用户端-v1.0.docx'
doc.save(output)
print('Done: ' + output)
